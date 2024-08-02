import sys
import os
import fitz  # PyMuPDF
import docx  # python-docx
import openpyxl  # openpyxl
from PyQt5.QtWidgets import (QApplication, QMainWindow, QFileDialog, QPushButton, QLabel,
                             QVBoxLayout, QWidget, QProgressBar, QTableWidget, QTableWidgetItem,
                             QLineEdit, QTextEdit, QComboBox, QMessageBox, QHBoxLayout, QHeaderView,
                             QMenuBar, QMenu)
from PyQt5.QtCore import QThread, pyqtSignal, QEvent, QDateTime
from PyQt5.QtGui import QTextCursor, QTextCharFormat, QBrush, QColor, QIcon
from concurrent.futures import ThreadPoolExecutor

class PhraseFinder(QMainWindow):
    def __init__(self):
        super().__init__()
        self.folder_path = ""
        self.search_thread = None
        self.initUI()

    def initUI(self):
        self.setWindowTitle('Advanced txtsearch')
        self.setGeometry(100, 100, 1200, 800)
        self.setWindowIcon(QIcon("path/to/icon.png"))

        # Create main layout
        main_layout = QVBoxLayout()

        # Create menu bar
        self.menuBar = QMenuBar(self)
        file_menu = QMenu('File', self)
        self.menuBar.addMenu(file_menu)
        exit_action = file_menu.addAction('Exit')
        exit_action.triggered.connect(self.close)
        self.setMenuBar(self.menuBar)

        # Top bar layout
        top_bar_layout = QHBoxLayout()

        # Widgets in Top Bar
        self.phrase_input = QLineEdit(self)
        self.phrase_input.setPlaceholderText('Enter phrase to search for')
        top_bar_layout.addWidget(self.phrase_input)

        self.extension_input = QComboBox(self)
        self.extension_input.addItems([".txt", ".pdf", ".docx", ".xlsx", "All Files"])
        top_bar_layout.addWidget(self.extension_input)

        self.folder_button = QPushButton('Select Folder', self)
        self.folder_button.clicked.connect(self.open_folder_dialog)
        top_bar_layout.addWidget(self.folder_button)

        self.search_button = QPushButton('Start Search', self)
        self.search_button.clicked.connect(self.start_search)
        top_bar_layout.addWidget(self.search_button)

        self.stop_button = QPushButton('Stop Search', self)
        self.stop_button.clicked.connect(self.stop_search)
        self.stop_button.setEnabled(False)
        top_bar_layout.addWidget(self.stop_button)

        self.save_button = QPushButton('Save Results', self)
        self.save_button.clicked.connect(self.save_results)
        top_bar_layout.addWidget(self.save_button)

        main_layout.addLayout(top_bar_layout)

        # Progress Bar
        self.progress_bar = QProgressBar(self)
        main_layout.addWidget(self.progress_bar)

        # Current File Label
        self.current_file_label = QLabel('Current File: None', self)
        main_layout.addWidget(self.current_file_label)

        # Results Table
        self.results_table = QTableWidget(self)
        self.results_table.setColumnCount(5)
        self.results_table.setHorizontalHeaderLabels(["File Path", "File Name", "Snippet", "Size", "Date Modified"])
        self.results_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.results_table.setSortingEnabled(True)
        self.results_table.cellDoubleClicked.connect(self.show_file_content)
        self.results_table.installEventFilter(self)
        main_layout.addWidget(self.results_table)

        # Status Label
        self.status_label = QLabel('Status: Idle', self)
        main_layout.addWidget(self.status_label)

        # File Content Text Edit
        self.file_content_text_edit = QTextEdit(self)
        self.file_content_text_edit.setReadOnly(True)
        main_layout.addWidget(self.file_content_text_edit)

        # Set Layout
        container = QWidget()
        container.setLayout(main_layout)
        self.setCentralWidget(container)

    def open_folder_dialog(self):
        self.folder_path = QFileDialog.getExistingDirectory(self, 'Select Folder')

    def start_search(self):
        if not self.folder_path:
            QMessageBox.warning(self, 'Warning', 'Please select a folder first.')
            return
        self.status_label.setText('Status: Searching...')
        self.results_table.setRowCount(0)
        self.progress_bar.setValue(0)
        self.search_phrase = self.phrase_input.text()
        self.search_thread = SearchThread(self.search_phrase, self.folder_path, self.extension_input.currentText())
        self.search_thread.update_progress.connect(self.update_progress)
        self.search_thread.update_results.connect(self.update_results)
        self.search_thread.update_current_file.connect(self.update_current_file)
        self.search_thread.finished.connect(self.search_finished)
        self.search_thread.start()
        self.search_button.setEnabled(False)
        self.stop_button.setEnabled(True)

    def stop_search(self):
        if self.search_thread:
            self.search_thread.requestInterruption()
            self.status_label.setText('Status: Stopping...')

    def search_finished(self):
        self.status_label.setText('Status: Done')
        self.search_button.setEnabled(True)
        self.stop_button.setEnabled(False)

    def update_progress(self, progress):
        self.progress_bar.setValue(progress)

    def update_results(self, file_path, file_name, snippet, size, date_modified):
        row_position = self.results_table.rowCount()
        self.results_table.insertRow(row_position)
        self.results_table.setItem(row_position, 0, QTableWidgetItem(file_path))
        self.results_table.setItem(row_position, 1, QTableWidgetItem(file_name))
        self.results_table.setItem(row_position, 2, QTableWidgetItem(snippet))
        self.results_table.setItem(row_position, 3, QTableWidgetItem(size))
        self.results_table.setItem(row_position, 4, QTableWidgetItem(date_modified))

    def update_current_file(self, file_path):
        self.current_file_label.setText(f'Current File: {file_path}')

    def show_file_content(self, row, column):
        file_path = self.results_table.item(row, 0).text()
        self.load_file_content(file_path)

    def eventFilter(self, source, event):
        if event.type() == QEvent.ToolTip and source is self.results_table:
            index = self.results_table.indexAt(event.pos())
            if index.isValid():
                file_path = self.results_table.item(index.row(), 0).text()
                self.load_file_content(file_path)
        return super().eventFilter(source, event)

    def load_file_content(self, file_path):
        try:
            content = self.read_file(file_path)
            self.display_file_content(content, self.search_phrase)
        except Exception as e:
            self.file_content_text_edit.setPlainText(f"Could not read file {file_path}: {e}")

    def read_file(self, file_path):
        ext = os.path.splitext(file_path)[1].lower()
        if ext == '.pdf':
            return self.read_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            return self.read_docx(file_path)
        elif ext in ['.xlsx', '.xls']:
            return self.read_xlsx(file_path)
        else:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            except Exception as e:
                return f"Could not read file {file_path}: {e}"

    def read_pdf(self, file_path):
        text = ""
        try:
            doc = fitz.open(file_path)
            for page in doc:
                text += page.get_text()
        except Exception as e:
            text = f"Could not read PDF file {file_path}: {e}"
        return text

    def read_docx(self, file_path):
        text = ""
        try:
            doc = docx.Document(file_path)
            text = "\n".join([p.text for p in doc.paragraphs])
        except Exception as e:
            text = f"Could not read DOCX file {file_path}: {e}"
        return text

    def read_xlsx(self, file_path):
        text = ""
        try:
            wb = openpyxl.load_workbook(file_path, data_only=True)
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                for row in ws.iter_rows(values_only=True):
                    text += " ".join([str(cell) for cell in row if cell is not None]) + "\n"
        except Exception as e:
            text = f"Could not read XLSX file {file_path}: {e}"
        return text

    def display_file_content(self, content, phrase):
        cursor = QTextCursor(self.file_content_text_edit.document())
        cursor.movePosition(QTextCursor.Start)
        cursor.select(QTextCursor.Document)
        cursor.setCharFormat(QTextCharFormat())  # Reset formatting
        self.file_content_text_edit.setPlainText(content)
        if phrase:
            cursor = QTextCursor(self.file_content_text_edit.document())
            cursor.setPosition(0)
            format = QTextCharFormat()
            format.setBackground(QBrush(QColor('yellow')))
            while not cursor.isNull() and not cursor.atEnd():
                cursor = self.file_content_text_edit.document().find(phrase, cursor)
                if not cursor.isNull():
                    cursor.mergeCharFormat(format)

    def save_results(self):
        file_path, _ = QFileDialog.getSaveFileName(self, 'Save Results', '', 'Text Files (*.txt);;All Files (*)')
        if file_path:
            try:
                with open(file_path, 'w', encoding='utf-8') as f:
                    for row in range(self.results_table.rowCount()):
                        file_path = self.results_table.item(row, 0).text()
                        file_name = self.results_table.item(row, 1).text()
                        snippet = self.results_table.item(row, 2).text()
                        size = self.results_table.item(row, 3).text()
                        date_modified = self.results_table.item(row, 4).text()
                        f.write(f"{file_path}\t{file_name}\t{snippet}\t{size}\t{date_modified}\n")
                QMessageBox.information(self, 'Success', 'Results saved successfully.')
            except Exception as e:
                QMessageBox.critical(self, 'Error', f'Could not save results: {e}')

class SearchThread(QThread):
    update_progress = pyqtSignal(int)
    update_results = pyqtSignal(str, str, str, str, str)
    update_current_file = pyqtSignal(str)

    def __init__(self, phrase, folder_path, extension):
        super().__init__()
        self.phrase = phrase
        self.folder_path = folder_path
        self.extension = extension
        self.interrupted = False

    def run(self):
        processed_files = 0
        total_files = sum([len(files) for r, d, files in os.walk(self.folder_path)])

        with ThreadPoolExecutor() as executor:
            future_to_file = {executor.submit(self.process_file, os.path.join(root, file)): os.path.join(root, file)
                              for root, dirs, files in os.walk(self.folder_path)
                              for file in files
                              if self.extension == "All Files" or file.endswith(self.extension)}

            for future in future_to_file:
                if self.isInterruptionRequested():
                    self.interrupted = True
                    break

                try:
                    future.result()  # Raises exception if the file processing failed
                except Exception as e:
                    print(f"Error processing file {future_to_file[future]}: {e}")

                processed_files += 1
                self.update_progress.emit(int((processed_files / total_files) * 100))

    def process_file(self, file_path):
        if self.isInterruptionRequested():
            return

        self.update_current_file.emit(file_path)

        try:
            content = self.read_file(file_path)
            if self.phrase.lower() in content.lower():
                snippet = self.get_snippet(content)
                size = os.path.getsize(file_path)
                date_modified = QDateTime.fromMSecsSinceEpoch(int(os.path.getmtime(file_path) * 1000)).toString()
                self.update_results.emit(file_path, os.path.basename(file_path), snippet, str(size), date_modified)
        except Exception as e:
            print(f"Could not process file {file_path}: {e}")

    def read_file(self, file_path):
        ext = os.path.splitext(file_path)[1].lower()
        if ext == '.pdf':
            return self.read_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            return self.read_docx(file_path)
        elif ext in ['.xlsx', '.xls']:
            return self.read_xlsx(file_path)
        else:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            except Exception as e:
                return f"Could not read file {file_path}: {e}"

    def read_pdf(self, file_path):
        text = ""
        try:
            doc = fitz.open(file_path)
            for page in doc:
                text += page.get_text()
        except Exception as e:
            text = f"Could not read PDF file {file_path}: {e}"
        return text

    def read_docx(self, file_path):
        text = ""
        try:
            doc = docx.Document(file_path)
            text = "\n".join([p.text for p in doc.paragraphs])
        except Exception as e:
            text = f"Could not read DOCX file {file_path}: {e}"
        return text

    def read_xlsx(self, file_path):
        text = ""
        try:
            wb = openpyxl.load_workbook(file_path, data_only=True)
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                for row in ws.iter_rows(values_only=True):
                    text += " ".join([str(cell) for cell in row if cell is not None]) + "\n"
        except Exception as e:
            text = f"Could not read XLSX file {file_path}: {e}"
        return text

    def get_snippet(self, content):
        # Basic snippet extraction; can be improved
        start_index = max(content.lower().find(self.phrase.lower()) - 30, 0)
        end_index = min(start_index + 60, len(content))
        return content[start_index:end_index].replace(self.phrase, f"<b>{self.phrase}</b>")

if __name__ == '__main__':
    app = QApplication(sys.argv)
    window = PhraseFinder()
    window.show()
    sys.exit(app.exec_())
